"""
task 2
pylint 10
"""
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
DOC1 = Document()
DOC1.add_paragraph("Hello Python")
DOC1.save("hello_python.docx")
print("Word файл 'hello_python.docx' успешно создан и заполнен текстом.")
DOC2 = Document("hello_python.docx")
bold_text = []
for paragraph in DOC2.paragraphs:
    for run in paragraph.runs:
        if run.bold:
            bold_text.append(run.text)
print("Жирный текст из файла 'hello_python.docx':", bold_text)
DOC3 = Document()
paragraph = DOC3.add_paragraph("Это абзац с измененным шрифтом и размером шрифта.")
run = paragraph.add_run()
font = run.font
font.name = 'Arial'
font.size = Pt(14)
paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
DOC3.save("styled_paragraph.docx")
print("Word файл 'styled_paragraph.docx' успешно создан с измененным стилем текста.")
